from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import PurePath

ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}
DEFAULT_MAX_BYTES = 1024 * 1024


@dataclass(frozen=True)
class ParsedDocument:
    name: str
    text: str


def safe_filename(filename: str | None) -> str:
    name = PurePath(filename or "upload.txt").name
    if not name or name in {".", ".."} or len(name) > 180:
        raise ValueError("invalid filename")
    return name


def parse_document(filename: str | None, content: bytes, max_bytes: int = DEFAULT_MAX_BYTES) -> ParsedDocument:
    name = safe_filename(filename)
    if len(content) > max_bytes:
        raise ValueError("file exceeds upload limit")
    suffix = PurePath(name).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError("unsupported file type; use TXT, PDF, or DOCX")
    try:
        if suffix == ".txt":
            text = content.decode("utf-8-sig")
        elif suffix == ".pdf":
            from pypdf import PdfReader
            text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(content)).pages)
        else:
            from docx import Document
            text = "\n".join(p.text for p in Document(io.BytesIO(content)).paragraphs)
    except ImportError as exc:
        raise ValueError(f"{suffix} parsing requires the optional documents install") from exc
    except Exception as exc:
        raise ValueError("could not parse document") from exc
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        raise ValueError("document contains no extractable text")
    return ParsedDocument(name=name, text=text)


def chunk_text(document: ParsedDocument, size: int = 90, overlap: int = 15) -> list[dict[str, str]]:
    if size <= overlap or size < 1:
        raise ValueError("chunk size must be greater than overlap")
    words = document.text.split()
    chunks = []
    step = size - overlap
    for start in range(0, len(words), step):
        part = " ".join(words[start : start + size])
        if part:
            chunks.append({"chunk_id": f"{document.name}#chunk-{len(chunks) + 1}", "document": document.name, "text": part})
        if start + size >= len(words):
            break
    return chunks
